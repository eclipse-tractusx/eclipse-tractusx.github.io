#********************************************************************************
# Copyright (c) 2026 Contributors to the Eclipse Foundation
#
# See the NOTICE file(s) distributed with this work for additional
# information regarding copyright ownership.
#
# This program and the accompanying materials are made available under the
# terms of the Apache License, Version 2.0 which is available at
# https://www.apache.org/licenses/LICENSE-2.0.
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS, WITHOUT
# WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied. See the
# License for the specific language governing permissions and limitations
# under the License.
#
# SPDX-License-Identifier: Apache-2.0
#*******************************************************************************/

#!/usr/bin/env python3
r"""
Generate a Word document from markdown files in a specified folder.
This script combines all markdown files from the specified folder
and converts them to a Word (.docx) document.

Usage:
    python generate_kit_word.py <folder_path> [output_file]
    
Examples:
    python generate_kit_word.py docs-kits/kit-template
    python generate_kit_word.py "docs-kits/kits/Business Partner Kit" business-partner-kit.docx
"""

import os
import sys
import argparse
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup
import re
import requests
from io import BytesIO
from PIL import Image as PILImage
from datetime import datetime

def collect_markdown_files(base_path):
    """
    Dynamically collect all markdown and MDX files from the directory.
    Files are sorted to maintain a logical order:
    1. README.md first
    2. Other files in alphabetical order by path
    """
    md_files = []
    
    # Find all .md and .mdx files recursively
    for pattern in ['**/*.md', '**/*.mdx']:
        md_files.extend(base_path.glob(pattern))
    
    # Remove duplicates and exclude README files
    md_files = [f for f in set(md_files) if f.stem.upper() != 'README']

    # Custom sort: changelog last, everything else alphabetically
    def sort_key(path):
        relative = path.relative_to(base_path)
        if path.stem.lower() == 'changelog':
            return (2, str(relative))
        return (1, str(relative))
    
    md_files.sort(key=sort_key)
    
    return md_files

def read_markdown_file(file_path):
    """Read markdown file content."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""

def _apply_font(style, font_name):
    """Set the given font and strip any theme-font overrides from a style's XML."""
    style.font.name = font_name
    rPr = style.font._element
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    for attr in [qn('w:asciiTheme'), qn('w:hAnsiTheme'), qn('w:cstheme'), qn('w:eastAsiaTheme')]:
        if attr in rFonts.attrib:
            del rFonts.attrib[attr]


def setup_document_styles(doc, body_font='Calibri'):
    """Set up custom styles for the Word document."""

    # Patch the document theme so major/minor fonts match, not Calibri/Calibri Light
    try:
        import re as _re
        theme_part = doc.part.part_related_by(
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme'
        )
        theme_xml = theme_part._blob.decode('utf-8')
        theme_xml = _re.sub(r'(<a:latin typeface=")[^"]*(")', rf'\g<1>{body_font}\2', theme_xml)
        theme_part._blob = theme_xml.encode('utf-8')
    except Exception:
        pass

    styles = doc.styles

    # Default Paragraph Font (catches any unstyled runs)
    try:
        _apply_font(styles['Default Paragraph Font'], body_font)
    except Exception:
        pass

    # Normal style
    try:
        normal_style = styles['Normal']
        _apply_font(normal_style, body_font)
        normal_style.font.size = Pt(11)
        normal_style.font.color.rgb = RGBColor(0, 0, 0)
    except KeyError:
        pass

    # Heading 1
    try:
        h1_style = styles['Heading 1']
        _apply_font(h1_style, body_font)
        h1_style.font.size = Pt(24)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(44, 62, 80)
    except KeyError:
        pass

    # Heading 2
    try:
        h2_style = styles['Heading 2']
        _apply_font(h2_style, body_font)
        h2_style.font.size = Pt(18)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(52, 73, 94)
    except KeyError:
        pass

    # Heading 3
    try:
        h3_style = styles['Heading 3']
        _apply_font(h3_style, body_font)
        h3_style.font.size = Pt(14)
        h3_style.font.bold = True
        h3_style.font.color.rgb = RGBColor(127, 140, 141)
    except KeyError:
        pass

    # Code style (keep monospaced)
    try:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
    except:
        pass

def _add_page_num_field(run, field):
    """Append a simple Word field (PAGE or NUMPAGES) to a run."""
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = field
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'end')
    run._r.append(fc1)
    run._r.append(instr)
    run._r.append(fc2)


def add_page_number_footer(doc, kit_title=None, template_url=None):
    """Add a 3-column footer: kit title LEFT | Page X of Y CENTER | source link RIGHT.
    Only configures section 0; all other sections inherit via is_linked_to_previous=True."""

    def _build_footer(footer):
        footer.is_linked_to_previous = False
        footer.paragraphs[0].clear()
        para = footer.paragraphs[0]

        # Two tab stops: center and right
        pPr = para._p.get_or_add_pPr()
        tabs_el = OxmlElement('w:tabs')
        center_tab = OxmlElement('w:tab')
        center_tab.set(qn('w:val'), 'center')
        center_tab.set(qn('w:pos'), '4675')
        right_tab = OxmlElement('w:tab')
        right_tab.set(qn('w:val'), 'right')
        right_tab.set(qn('w:pos'), '9350')
        tabs_el.append(center_tab)
        tabs_el.append(right_tab)
        pPr.append(tabs_el)

        # LEFT: kit title
        left_run = para.add_run(kit_title if kit_title else '')
        left_run.font.size = Pt(9)
        left_run.font.color.rgb = RGBColor(100, 100, 100)

        # CENTER: Page X of Y
        para.add_run('\t')
        page_run = para.add_run()
        page_run.font.size = Pt(9)
        page_run.font.color.rgb = RGBColor(100, 100, 100)
        _add_page_num_field(page_run, 'PAGE')
        page_run.add_text(' of ')
        _add_page_num_field(page_run, 'NUMPAGES')

        # RIGHT: source link
        para.add_run('\t')
        if template_url:
            add_hyperlink(para, 'KIT Template Source', template_url)
        else:
            right_run = para.add_run('Eclipse Tractus-X')
            right_run.font.size = Pt(8)
            right_run.font.color.rgb = RGBColor(100, 100, 100)

    for i, section in enumerate(doc.sections):
        if i == 0:
            # Build the footer on section 0
            _build_footer(section.footer)
            # Also apply to first-page footer of section 0 so page 1 has it too
            if section.different_first_page_header_footer:
                _build_footer(section.first_page_footer)
        else:
            # All other sections inherit from the previous section
            section.footer.is_linked_to_previous = True

def add_toc(doc):
    """Insert a Word Table of Contents field on its own page."""
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Word TOC fields require each fldChar and instrText in its own <w:r>
    paragraph = doc.add_paragraph()
    p = paragraph._p

    def _make_run(*children):
        r = OxmlElement('w:r')
        for child in children:
            r.append(child)
        return r

    # <w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>
    fc_begin = OxmlElement('w:fldChar')
    fc_begin.set(qn('w:fldCharType'), 'begin')
    fc_begin.set(qn('w:dirty'), 'true')
    p.append(_make_run(fc_begin))

    # <w:r><w:instrText> TOC \o "1-3" \h \z \u </w:instrText></w:r>
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    p.append(_make_run(instr))

    # <w:r><w:fldChar w:fldCharType="separate"/></w:r>
    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    p.append(_make_run(fc_sep))

    # Placeholder text shown before Word updates the field
    placeholder_r = OxmlElement('w:r')
    placeholder_t = OxmlElement('w:t')
    placeholder_t.text = 'Right-click here and select "Update Field" to generate the Table of Contents.'
    placeholder_r.append(placeholder_t)
    p.append(placeholder_r)

    # <w:r><w:fldChar w:fldCharType="end"/></w:r>
    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    p.append(_make_run(fc_end))

    note = doc.add_paragraph()
    note_run = note.add_run('Tip: In Word, press Ctrl+A then F9 to update all fields including this TOC.')
    note_run.font.size = Pt(9)
    note_run.font.italic = True
    note_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

def add_section_title_page(doc, title, file_label=None):
    """Add a new Word section, set its running header, then render a compact section title bar."""
    new_section = doc.add_section()
    # Disable different-first-page so the header shows on every page of this section
    new_section.different_first_page_header_footer = False
    new_section.header.is_linked_to_previous = False

    header = new_section.header
    # Clear any existing paragraphs
    for p in header.paragraphs:
        p.clear()

    if not header.paragraphs:
        header_para = header.add_paragraph()
    else:
        header_para = header.paragraphs[0]

    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.add_run(file_label if file_label else title)
    header_run.font.size = Pt(9)
    header_run.font.italic = True
    header_run.font.color.rgb = RGBColor(100, 100, 100)

    # Thin rule under header text
    pPr = header_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Compact title bar — no spacers, content follows immediately
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.paragraph_format.space_before = Pt(6)
    title_para.paragraph_format.space_after = Pt(4)

    # Thin decorative rule below title
    rule_para = doc.add_paragraph()
    rule_para.paragraph_format.space_before = Pt(0)
    rule_para.paragraph_format.space_after = Pt(12)
    pPr2 = rule_para._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    bot2 = OxmlElement('w:bottom')
    bot2.set(qn('w:val'), 'single')
    bot2.set(qn('w:sz'), '6')
    bot2.set(qn('w:space'), '1')
    bot2.set(qn('w:color'), '2C3E50')
    pBdr2.append(bot2)
    pPr2.append(pBdr2)

def extract_frontmatter_title(content):
    """Extract the title field from YAML frontmatter, or return None."""
    match = re.match(r'^---\s*\n(.*?)\n---\s*\n', content, flags=re.DOTALL)
    if match:
        for line in match.group(1).splitlines():
            m = re.match(r"^title:\s*[\'\"]?(.*?)[\'\"]?\s*$", line)
            if m:
                return m.group(1).strip().strip("\"'")
    return None

def clean_markdown_content(content):
    """Clean markdown content by removing frontmatter, comments, and MDX/JSX specifics."""
    # Remove frontmatter (YAML between ---)
    content = re.sub(r'^---\s*\n.*?\n---\s*\n', '', content, flags=re.DOTALL | re.MULTILINE)
    
    # Remove HTML comments (including KIT LOGO blocks and REUSE blocks)
    content = re.sub(r'<!--.*?-->', '', content, flags=re.DOTALL)
    
    # Remove JSX/MDX import statements  (e.g. import Kit3DLogo from '...')
    content = re.sub(r'^import\s+\w+\s+from\s+[^\n]+\n?', '', content, flags=re.MULTILINE)
    
    # Remove JSX self-closing and open/close component tags.
    # Use DOTALL + non-greedy so attributes containing '<' (e.g. kitId="<kit-id>") are handled.
    content = re.sub(r'<[A-Z]\w+\b.*?/>', '', content, flags=re.DOTALL)   # self-closing
    content = re.sub(r'<[A-Z]\w+\b.*?>', '', content, flags=re.DOTALL)    # opening tag
    content = re.sub(r'</[A-Z]\w+>', '', content)                           # closing tag
    
    # Remove docusaurus-specific syntax like :::info, :::tip, etc.
    content = re.sub(r':::(\w+)\s*(.*?)\n', r'**\1**: \2\n', content)
    content = re.sub(r':::', '', content)
    
    # Handle mermaid diagrams - convert to special marker
    def replace_mermaid(match):
        mermaid_code = match.group(1).strip()
        return f'\n\n**[MERMAID_DIAGRAM_START]**\n```\n{mermaid_code}\n```\n**[MERMAID_DIAGRAM_END]**\n\n'
    
    content = re.sub(r'```mermaid\s*(.*?)```', replace_mermaid, content, flags=re.DOTALL)
    
    return content.strip()

def markdown_to_html(md_content):
    """Convert markdown content to HTML."""
    md = markdown.Markdown(extensions=[
        'extra',
        'tables',
        'fenced_code',
        'nl2br'
    ])
    return md.convert(md_content)

def add_image_to_docx(doc, image_path, base_path, max_width=6.0):
    """Add an image to the Word document."""
    try:
        # Get repo root (two levels up from base_path which is in docs-kits/)
        repo_root = base_path
        while repo_root.parent != repo_root and not (repo_root / 'static').exists():
            repo_root = repo_root.parent
        
        # Resolve image path
        if image_path.startswith('http://') or image_path.startswith('https://'):
            # Download image from URL
            response = requests.get(image_path, timeout=10)
            if response.status_code == 200:
                image_stream = BytesIO(response.content)
            else:
                print(f"Failed to download image: {image_path}")
                return
        else:
            # Handle relative paths
            if image_path.startswith('/'):
                # Absolute path from repo root
                img_file = repo_root / image_path.lstrip('/')
            elif image_path.startswith('../'):
                # Relative path - resolve from the file's directory
                img_file = (base_path / image_path).resolve()
            else:
                # Relative path from base
                img_file = base_path / image_path
            
            if not img_file.exists():
                print(f"Image not found: {img_file}")
                return
            
            image_stream = str(img_file)
        
        # Add image with size constraints
        try:
            # Check if it's an SVG file
            if isinstance(image_stream, str) and image_stream.lower().endswith('.svg'):
                # For SVG files, add a placeholder text instead
                p = doc.add_paragraph(f"[SVG Image: {Path(image_stream).name}]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.style.font.italic = True
                p.style.font.size = Pt(10)
                p.style.font.color.rgb = RGBColor(100, 100, 100)
                print(f"Note: SVG image converted to placeholder: {Path(image_stream).name}")
                return
            
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            
            # Get image dimensions to maintain aspect ratio
            if isinstance(image_stream, str):
                img = PILImage.open(image_stream)
            else:
                img = PILImage.open(image_stream)
            
            width, height = img.size
            aspect_ratio = height / width
            
            # Calculate display size
            display_width = min(max_width, width / 96)  # Convert pixels to inches (96 DPI)
            display_height = display_width * aspect_ratio
            
            if isinstance(image_stream, str):
                run.add_picture(image_stream, width=Inches(display_width))
            else:
                run.add_picture(image_stream, width=Inches(display_width))
            
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as img_err:
            print(f"Error adding image to document: {img_err}")
    
    except Exception as e:
        print(f"Error processing image {image_path}: {e}")

def process_mermaid_diagram(doc, mermaid_code):
    """Convert mermaid diagram to PNG using Python and add to the document."""
    # Clean the markers from the code
    mermaid_code = re.sub(r'\[MERMAID_DIAGRAM_START\]|\[MERMAID_DIAGRAM_END\]', '', mermaid_code)
    mermaid_code = mermaid_code.strip()
    
    # Try to render as PNG using mermaid.ink API
    import tempfile
    import base64
    
    diagram_rendered = False
    
    try:
        # Use mermaid.ink API with base64 encoding
        encoded_mermaid = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
        mermaid_url = f"https://mermaid.ink/img/{encoded_mermaid}"
        
        # Download the rendered image
        response = requests.get(mermaid_url, timeout=15)
        
        if response.status_code == 200 and 'image' in response.headers.get('Content-Type', ''):
            # Save to temporary file
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                tmp_file.write(response.content)
                png_path = tmp_file.name
            
            # Add diagram title
            title_para = doc.add_paragraph()
            title_run = title_para.add_run('📊 Diagram')
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = RGBColor(44, 62, 80)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_before = Pt(12)
            title_para.paragraph_format.space_after = Pt(6)
            
            # Add the rendered image
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            
            # Get image dimensions and scale appropriately
            from PIL import Image
            img = Image.open(png_path)
            width, height = img.size
            aspect_ratio = height / width
            
            # Scale to fit page (max 6 inches wide)
            max_width = 6.0
            display_width = min(max_width, width / 96)  # 96 DPI
            
            run.add_picture(png_path, width=Inches(display_width))
            
            # Clean up temp file
            Path(png_path).unlink(missing_ok=True)
            
            print(f"  Rendered mermaid diagram from mermaid.ink")
            diagram_rendered = True
        else:
            print(f"  Warning: mermaid.ink API returned status {response.status_code}")
            
    except Exception as e:
        print(f"  Warning: Could not render mermaid diagram ({e})")
    
    # Always add the source code below the diagram (or standalone if rendering failed)
    if diagram_rendered:
        # Add a subtitle indicating this is the source code
        code_title = doc.add_paragraph()
        code_title_run = code_title.add_run('Source Code:')
        code_title_run.font.size = Pt(9)
        code_title_run.italic = True
        code_title_run.font.color.rgb = RGBColor(100, 100, 100)
        code_title.paragraph_format.space_before = Pt(6)
        code_title.paragraph_format.space_after = Pt(3)
    else:
        # If rendering failed, show a clear title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('📊 Diagram (Code)')
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = RGBColor(44, 62, 80)
        title_para.paragraph_format.space_before = Pt(12)
        title_para.paragraph_format.space_after = Pt(6)
    
    # Add diagram code in a shaded box
    code_para = doc.add_paragraph()
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_para.paragraph_format.right_indent = Inches(0.5)
    code_para.paragraph_format.space_after = Pt(12)
    
    code_run = code_para.add_run(mermaid_code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    code_run.font.color.rgb = RGBColor(44, 62, 80)
    
    # Add shading to the paragraph
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F8F9FA')
    code_para._p.get_or_add_pPr().append(shading_elm)

def add_hint_box(doc, text, is_todo=False):
    """
    Render a blockquote as a styled callout box.
    - TODO blockquotes: amber/yellow background with a ✏ icon label.
    - Regular blockquotes: light-blue background, italic.
    """
    if is_todo:
        # Strip leading "TODO:" or "TODO" from the text for the body
        body = re.sub(r'^TODO[:\s]*', '', text, flags=re.IGNORECASE).strip()
        fill_color = 'FFF3CD'   # amber-yellow
        border_color = 'FFCA28'
        label = '✏  TODO'
        label_color = RGBColor(0x85, 0x64, 0x04)
        body_color  = RGBColor(0x66, 0x4D, 0x03)
    else:
        body = text
        fill_color = 'E8F4FD'   # light blue
        border_color = '5DADE2'
        label = None
        label_color = None
        body_color  = RGBColor(0x1A, 0x5E, 0x7A)

    def _shade_para(p, fill):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        p._p.get_or_add_pPr().append(shd)

    def _border_para(p, color):
        """Add a thick left border to a paragraph to simulate a callout bar."""
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'thick')
        left.set(qn('w:sz'), '24')       # 3pt
        left.set(qn('w:space'), '8')
        left.set(qn('w:color'), color)
        pBdr.append(left)
        pPr.append(pBdr)

    indent = Inches(0.25)

    if label:
        label_p = doc.add_paragraph()
        label_p.paragraph_format.left_indent  = indent
        label_p.paragraph_format.space_before = Pt(8)
        label_p.paragraph_format.space_after  = Pt(0)
        label_run = label_p.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(10)
        label_run.font.color.rgb = label_color
        _shade_para(label_p, fill_color)
        _border_para(label_p, border_color)

    body_p = doc.add_paragraph()
    body_p.paragraph_format.left_indent  = indent
    body_p.paragraph_format.space_before = Pt(0) if label else Pt(8)
    body_p.paragraph_format.space_after  = Pt(8)
    body_run = body_p.add_run(body)
    body_run.font.size = Pt(10)
    body_run.font.italic = not is_todo
    body_run.font.color.rgb = body_color
    _shade_para(body_p, fill_color)
    _border_para(body_p, border_color)


def add_html_to_docx(doc, html_content, base_path):
    """Parse HTML and add content to Word document."""
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Track if we're inside a mermaid diagram
    in_mermaid = False
    mermaid_code = []
    
    for element in soup.children:
        if element.name is None:
            continue
        
        # Check for mermaid diagram markers
        element_text = element.get_text().strip()
        
        if '[MERMAID_DIAGRAM_START]' in element_text:
            in_mermaid = True
            mermaid_code = []
            continue
        elif '[MERMAID_DIAGRAM_END]' in element_text:
            if in_mermaid and mermaid_code:
                # Process the collected mermaid code
                full_code = '\n'.join(mermaid_code)
                process_mermaid_diagram(doc, full_code)
            in_mermaid = False
            mermaid_code = []
            continue
        
        # If we're collecting mermaid code
        if in_mermaid:
            if element.name == 'pre':
                code_text = element.get_text().strip()
                mermaid_code.append(code_text)
            continue
            
        # Headings
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            text = element.get_text().strip()
            if text:
                if level == 1:
                    doc.add_page_break()
                doc.add_heading(text, level=min(level, 3))
        
        # Images
        elif element.name == 'img':
            img_src = element.get('src', '')
            img_alt = element.get('alt', '')
            
            if img_src:
                # Add alt text as caption if available
                if img_alt:
                    caption = doc.add_paragraph(img_alt)
                    caption.style.font.italic = True
                    caption.style.font.size = Pt(9)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                add_image_to_docx(doc, img_src, base_path)
                doc.add_paragraph()  # Add spacing after image
        
        # Paragraphs
        elif element.name == 'p':
            # Check for images within paragraphs
            imgs = element.find_all('img')
            if imgs:
                for img in imgs:
                    img_src = img.get('src', '')
                    img_alt = img.get('alt', '')
                    if img_src:
                        if img_alt:
                            caption = doc.add_paragraph(img_alt)
                            caption.style.font.italic = True
                            caption.style.font.size = Pt(9)
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        add_image_to_docx(doc, img_src, base_path)
            else:
                text = element.get_text().strip()
                if text:
                    p = doc.add_paragraph(text)
        
        # Code blocks
        elif element.name == 'pre':
            code_text = element.get_text().strip()
            if code_text:
                try:
                    p = doc.add_paragraph(code_text, style='Code')
                except:
                    p = doc.add_paragraph(code_text)
                    p.style.font.name = 'Courier New'
                    p.style.font.size = Pt(9)
        
        # Blockquotes
        elif element.name == 'blockquote':
            text = element.get_text().strip()
            if text:
                is_todo = text.upper().startswith('TODO')
                add_hint_box(doc, text, is_todo=is_todo)
        
        # Lists
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                text = li.get_text().strip()
                if text:
                    doc.add_paragraph(text, style='List Bullet' if element.name == 'ul' else 'List Number')
        
        # Tables
        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                # Get dimensions
                num_rows = len(rows)
                num_cols = max(len(row.find_all(['th', 'td'])) for row in rows)
                
                if num_cols > 0:
                    # Create table
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    # Fill table
                    for i, row in enumerate(rows):
                        cells = row.find_all(['th', 'td'])
                        for j, cell in enumerate(cells):
                            if j < num_cols:
                                table.rows[i].cells[j].text = cell.get_text().strip()
                                # Bold header cells
                                if cell.name == 'th':
                                    for paragraph in table.rows[i].cells[j].paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
        
        # Horizontal rules
        elif element.name == 'hr':
            doc.add_paragraph('_' * 80)

def process_markdown_file(doc, file_path, base_path, source_path):
    """Process a single markdown file and add to document."""
    print(f"Processing: {file_path.name}")
    
    # Read raw content (before cleaning) to extract frontmatter title
    raw_content = read_markdown_file(file_path)
    if not raw_content:
        return False
    
    # Extract title from frontmatter
    section_title = extract_frontmatter_title(raw_content)
    if not section_title:
        # Fall back to a prettified filename
        section_title = file_path.stem.replace('-', ' ').replace('_', ' ').title()
    
    # Clean content (strips frontmatter, comments, JSX, imports, etc.)
    content = clean_markdown_content(raw_content)
    
    # Check if content has meaningful text
    stripped_content = re.sub(r'<[^>]+>', '', content)
    stripped_content = re.sub(r'\s+', ' ', stripped_content).strip()
    
    if len(stripped_content) < 50:
        print(f"  Skipping {file_path.name} - insufficient content ({len(stripped_content)} chars)")
        return False
    
    # Add a prominent section title page for every file
    # Use the relative path (e.g. adoption-view/adoption-view.md) as the running page header
    try:
        rel_label = str(file_path.relative_to(source_path))
    except ValueError:
        rel_label = file_path.name
    add_section_title_page(doc, section_title, file_label=rel_label)
    
    # Convert to HTML and add to document
    html_content = markdown_to_html(content)
    add_html_to_docx(doc, html_content, file_path.parent)
    
    doc.add_paragraph()  # trailing spacing
    return True

def extract_readme_intro(source_path):
    """
    Extract the first descriptive paragraph and any requirements URL from README.md.
    Returns (description, requirements_url) — either may be None.
    """
    readme = source_path / 'README.md'
    if not readme.exists():
        return None, None

    text = readme.read_text(encoding='utf-8')
    description = None
    req_url = None

    for line in text.splitlines():
        stripped = line.strip()
        # Skip headings, blank lines, blockquotes, code fences, badge lines
        if not stripped or stripped.startswith('#') or stripped.startswith('>') \
                or stripped.startswith('```') or stripped.startswith('!'):
            continue
        # Pick up a requirements URL (line containing the kit-lifecycle anchor)
        if 'kit-lifecycle' in stripped or ('requirements' in stripped.lower() and 'http' in stripped):
            url_match = re.search(r'https?://\S+', stripped)
            if url_match and req_url is None:
                req_url = url_match.group(0).rstrip(')')
            continue
        # First plain-text line is the description
        if description is None and not stripped.startswith('-') and not stripped.startswith('|'):
            # Strip inline markdown (bold, italic, backticks, links)
            clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', stripped)
            clean = re.sub(r'[*_`]', '', clean)
            if len(clean) > 20:
                description = clean

    return description, req_url


def extract_readme_title(source_path):
    """Extract the first # heading from README.md as the document title."""
    readme = source_path / 'README.md'
    if not readme.exists():
        return None
    for line in readme.read_text(encoding='utf-8').splitlines():
        if line.startswith('# '):
            return line[2:].strip()
    return None


def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '2980B9')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(color)
    rPr.append(u)
    r = OxmlElement('w:r')
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)
    return hyperlink

def generate_word_doc(source_path, output_path, body_font='Calibri'):
    """
    Generate a Word document from markdown files in the specified folder.
    
    Args:
        source_path: Path to the directory containing markdown files
        output_path: Path where the Word document should be saved
        body_font: Font name to use for body and heading text
    """
    print(f"Starting Word document generation from: {source_path}")
    
    # Create document
    doc = Document()
    
    # Setup styles
    setup_document_styles(doc, body_font=body_font)

    folder_name = source_path.name
    readme_title = extract_readme_title(source_path) or f'{folder_name} KIT Template'
    template_url = (
        f'https://github.com/eclipse-tractusx/eclipse-tractusx.github.io'
        f'/tree/main/docs-kits/kit-template/{folder_name}'
    )
    readme_desc, req_url = extract_readme_intro(source_path)

    # First section: first-page header shows centered doc title; regular header stays empty
    first_section = doc.sections[0]
    first_section.different_first_page_header_footer = True

    # Regular header (page 2+ of TOC section): empty
    first_section.header.is_linked_to_previous = False
    for p in first_section.header.paragraphs:
        p.clear()

    # First-page header: centered "Eclipse Tractus-X <title>"
    fph = first_section.first_page_header
    fph.is_linked_to_previous = False
    for p in fph.paragraphs:
        p.clear()
    fph_para = fph.paragraphs[0] if fph.paragraphs else fph.add_paragraph()
    fph_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fph_run = fph_para.add_run(f'Eclipse Tractus-X {readme_title}'.upper())
    fph_run.font.size = Pt(11)
    fph_run.font.bold = True
    fph_run.font.color.rgb = RGBColor(44, 62, 80)

    # Description + requirements link at the very start of the body
    if readme_desc or req_url:
        if readme_desc:
            desc_para = doc.add_paragraph(readme_desc)
            desc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            desc_r = desc_para.runs[0]
            desc_r.font.size = Pt(10)
            desc_r.font.italic = True
            desc_r.font.color.rgb = RGBColor(80, 80, 80)
        if req_url:
            req_para = doc.add_paragraph()
            req_label = req_para.add_run('Requirements can be found here: ')
            req_label.font.size = Pt(9)
            req_label.font.bold = True
            req_label.font.color.rgb = RGBColor(80, 80, 80)
            add_hyperlink(req_para, req_url, req_url)

    # Disclaimer
    disclaimer_para = doc.add_paragraph()
    disclaimer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    disclaimer_para.paragraph_format.space_before = Pt(14)
    disclaimer_para.paragraph_format.space_after = Pt(14)
    disclaimer_label = disclaimer_para.add_run('⚠ Disclaimer: ')
    disclaimer_label.font.size = Pt(9)
    disclaimer_label.font.bold = True
    disclaimer_label.font.color.rgb = RGBColor(133, 77, 14)
    disclaimer_text = disclaimer_para.add_run(
        'This Word document is a reference template only and will not be published as part of the KIT. '
        'In order to publish a KIT, the content of this document must be manually copied into the KIT template repository and adapted to fit the KIT structure and guidelines.'
        'The Table of Contents is not part of the KIT structure and should be removed before publication, it is only relevant for this word document.'
    )
    disclaimer_text.font.size = Pt(9)
    disclaimer_text.font.italic = True
    disclaimer_text.font.color.rgb = RGBColor(133, 77, 14)

    # Table of Contents fills the rest of page 1
    add_toc(doc)
    
    # Collect all markdown files
    md_files = collect_markdown_files(source_path)
    print(f"Found {len(md_files)} markdown files to process")
    
    if not md_files:
        print("Error: No markdown files found!")
        sys.exit(1)
    
    # Process each file — every file gets its own section title page
    files_added = 0
    for md_file in md_files:
        was_added = process_markdown_file(doc, md_file, md_file.parent, source_path)
        if was_added:
            files_added += 1
    
    # Add 3-column footer (kit title | page# | source link) to all sections
    add_page_number_footer(doc, kit_title=readme_title, template_url=template_url)
    
    # Save document
    print("Saving Word document...")
    doc.save(output_path)
    
    print(f"Word document generated successfully: {output_path}")
    print(f"File size: {os.path.getsize(output_path) / 1024:.2f} KB")

def main():
    """Main function."""
    parser = argparse.ArgumentParser(
        description='Generate a Word document from markdown files in a specified folder.',
        epilog='Example: python generate_kit_word.py docs-kits/kit-template'
    )
    parser.add_argument(
        'folder',
        nargs='?',
        help='Path to the folder containing markdown files (relative to repo root or absolute path)'
    )
    parser.add_argument(
        'output',
        nargs='?',
        help='Output file path (optional, defaults to <folder-name>.docx in repo root)'
    )
    parser.add_argument(
        '--font',
        default='Calibri',
        help='Body font to use (default: Calibri)'
    )

    args = parser.parse_args()
    
    # Get the repository root (assuming script is in scripts/ directory)
    script_dir = Path(__file__).parent
    repo_root = script_dir.parent
    
    # Determine source path
    if args.folder:
        source_path = Path(args.folder)
        # If relative path, make it relative to repo root
        if not source_path.is_absolute():
            source_path = repo_root / source_path
    else:
        # Default to kit-template for backward compatibility
        source_path = repo_root / "docs-kits" / "kit-template"
        print("No folder specified, using default: docs-kits/kit-template")
    
    # Check if source directory exists
    if not source_path.exists():
        print(f"Error: Source directory not found at {source_path}")
        sys.exit(1)
    
    if not source_path.is_dir():
        print(f"Error: {source_path} is not a directory")
        sys.exit(1)
    
    # Determine output path
    if args.output:
        output_path = Path(args.output)
        if not output_path.is_absolute():
            output_path = repo_root / output_path
    else:
        # Generate output filename from folder name
        folder_name = source_path.name.lower().replace(' ', '-')
        output_path = repo_root / f"{folder_name}.docx"
    
    # Generate the Word document
    try:
        generate_word_doc(source_path, output_path, body_font=args.font)
    except Exception as e:
        print(f"Error generating Word document: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()
